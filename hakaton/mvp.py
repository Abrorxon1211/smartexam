import tkinter as tk
from tkinter import messagebox
import random
from docx import Document
import threading
from time import sleep
from PIL import Image, ImageTk 
import pygame
import time
import os
import requests
import sounddevice as sd
import tempfile
import scipy.io.wavfile

# Fayl yo‘llari
word_file_path = os.path.join(os.path.expanduser("~"), "Desktop", "questions.docx")

IAM_TOKEN = "your_token"
FOLDER_ID = "your_id"


# mikrafon


def tts_stt(son):
    count = 0

    # 🎙️ Mikrofon sozlamalari
    duration = 5  # sekund
    samplerate = 16000
    channels = 1

    print("🔴 Gapiring...")

    # ⏺️ Ovoz yozish
    recording = sd.rec(int(duration * samplerate), samplerate=samplerate, channels=channels, dtype='int16')
    sd.wait()
    print("✅ Yozuv tugadi!")

    # 💾 Vaqtinchalik .wav fayl yaratish
    with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_wav:
        scipy.io.wavfile.write(temp_wav.name, samplerate, recording)
        audio_path = temp_wav.name

    # 📤 Yandex STT API ga so'rov
    with open(audio_path, "rb") as f:
        audio_data = f.read()

    headers = {
        "Authorization": f"Api-Key {IAM_TOKEN}"
    }

    params = {
        "folderId": FOLDER_ID,
        "lang": "uz-UZ",
        "sampleRateHertz": 16000,
        "format": "lpcm",
        "profanityFilter": "false"
    }

    response = requests.post(
        "https://stt.api.cloud.yandex.net/speech/v1/stt:recognize",
        headers=headers,
        params=params,
        data=audio_data
    )

    # 📋 Natijani ko‘rsatish va faylga yozish
    output_file = "response.txt"

    if response.status_code == 200:
        result = response.json()
        if "result" in result:
            recognized_text = result["result"]
            print("🗣️ Aniqlangan matn:", recognized_text)

            # 📝 Matnni faylga yozish
            with open(output_file, "a", encoding="utf-8") as f:
                f.write(f"({son}-savol javobi) " + recognized_text + '\n')
            print(f"📄 Matn '{output_file}' fayliga yozildi.")

            # 💬 Javobni GUIga yozish
            answers_box.insert(tk.END, recognized_text + "\n")
            answers_box.update()  # GUI yangilansin

        else:
            print("❌ Matn aniqlanmadi:", result)
    else:
        print("❌ Xatolik:", response.status_code, response.text)
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(response.text)
        print(f"📄 Xatolik matni '{output_file}' fayliga yozildi.")


def speak(a):
    text = a

    params = {
        "text": text,
        "lang": "uz-UZ",
        "voice": "nigora",
        "folderId": FOLDER_ID,
        "format": "mp3"
    }

    headers = {
        "Authorization": f"Api-Key {IAM_TOKEN}"
    }

    response = requests.post(
        "https://tts.api.cloud.yandex.net/speech/v1/tts:synthesize",
        headers=headers,
        params=params
    )

    if response.status_code == 200:
        file_path = "output.mp3"  # 🔁 FAYL NOMI
        with open(file_path, "wb") as f:
            f.write(response.content)
        print("✅ Audio saqlandi:", file_path)

        # 🔊 Pygame yordamida audio ijro etish
        pygame.mixer.init()
        pygame.mixer.music.load(file_path)
        pygame.mixer.music.play()

        # ⏳ Audio tugaguncha kutish
        while pygame.mixer.music.get_busy():
            time.sleep(0.1)

        # 🎯 Audio ijro etilgach tozalash
        pygame.mixer.quit()
        os.remove(file_path)
        print("🗑️ Audio fayl o‘chirildi:", file_path)

    else:
        print("❌ Xatolik:", response.status_code, response.text)


def get_random_questions(docx_path, count=3):
    doc = Document(docx_path)
    questions = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return random.sample(questions, min(count, len(questions)))


listening = False
selected_questions = []


def start_questions():
    global listening, selected_questions
    listening = True
    try:
        selected_questions = get_random_questions(word_file_path)
        questions_box.delete("1.0", tk.END)

        for i, question in enumerate(selected_questions, 1):
            question_str = f"({i}-savol) {question}\n"
            with open('response.txt', "a", encoding="utf-8") as f:
                f.write(question_str)

            # Savollarni GUIga qo‘shamiz
            questions_box.insert(tk.END, question_str)
            questions_box.update()  # GUI yangilansin

            # Savolni ovozli o‘qish
            speak(question)
            sleep(2)  # Keyingi savolga o‘tishdan oldin kutish

        # Javoblarni yozib olishga tayyorlanadi
        speak("Javoblaringiz yozib olinadi. Gapiring.")

        for i in range(1, 4):
            threading.Thread(target=tts_stt, args=(i,), daemon=True).start()
            sleep(10)
#####################################################################################################
        # API url va headers
        url = "https://gpt-4o.p.rapidapi.com/chat/completions"
        headers = {
            'x-rapidapi-key': "your_api_key",
            'x-rapidapi-host': "gpt-4o.p.rapidapi.com",
            'Content-Type': "application/json"
        }

        # Foydalanuvchidan savol olish

        with open("response.txt", "r", encoding="utf-8") as f:
            user_question = f.read()

            # Prompt qo'shish (kontekst qo'shish)
        prompt = """
            Sen — savollarga o‘quvchilarning javoblarini tekshiradigan tajribali o‘qituvchi yordamchisan.
        Sening vazifang — javobning to‘g‘riligi, to‘liqligi va mantiqiyligini baholash, xatolik yoki noaniqliklarni ko‘rsatish, 10 ballik tizim asosida baho qo‘yish.

        Quyidagi shablonga amal qil:

        Baho: [0 dan 10 gacha ball]

        Kontekst:
        id savol = savol matni
        id savol javobi = o'quvchi javobi matni
            """

        # Payload yaratish (savolni va promptni birlashtirish)
        payload = {
            "model": "gpt-4o",
            "messages": [
                {
                    "role": "system",  # Bu rolda tizimdan kelgan yordamchi xabar
                    "content": prompt
                },
                {
                    "role": "user",  # Bu foydalanuvchidan kelgan so'rov
                    "content": user_question
                }
            ]
        }

        # API'ga so'rov yuborish
        response = requests.post(url, json=payload, headers=headers)

        # Javobni chiqarish
        if response.status_code == 200:
            print("Javob:", response.json()['choices'][0]['message']['content'])
        else:
            print("Xatolik yuz berdi:", response.status_code)

    except Exception as e:
        messagebox.showerror("Xatolik", f"Word fayl ochilmadi: {e}")


def stop_listening():
    global listening
    listening = False
    speak("Dastur to‘xtatildi.")
    messagebox.showinfo("To‘xtatildi", "Yozib olish to‘xtatildi.")


def evaluate_answers():
    try:
        user_answers = answers_box.get("1.0", tk.END).strip().lower()
        doc = Document(answer_file_path)
        correct_answers = [p.text.strip().lower() for p in doc.paragraphs if p.text.strip()]

        if not selected_questions:
            messagebox.showwarning("Ogohlantirish", "Avval savollarni boshlang.")
            return

        score = 0
        max_score = len(selected_questions)

        for question in selected_questions:
            for correct in correct_answers:
                if correct in user_answers and any(q in correct for q in question.split()):
                    score += 1
                    break

        percentage = round((score / max_score) * 100)
        messagebox.showinfo("Natija", f"Siz {score} ta savolga to‘g‘ri javob berdingiz.\nBahoyingiz: {percentage}%")

    except Exception as e:
        messagebox.showerror("Xatolik", f"Baholashda xatolik: {e}")


# === GUI interfeys ===
window = tk.Tk()
window.title("SmartExam")
window.geometry("600x620")
window.config(bg="#f0f0f0")

# ✅ LOGOTIP eng yuqorida
try:
    logo_image = Image.open("logo.png")
    logo_image = logo_image.resize((400, 300))
    logo_photo = ImageTk.PhotoImage(logo_image)
    logo_label = tk.Label(window, image=logo_photo, bg="#f0f0f0")
    logo_label.image = logo_photo
    logo_label.pack(pady=5)
except Exception as e:
    print(f"Logotip yuklanmadi: {e}")

# Tugmalar
start_btn = tk.Button(window, text="Boshlash", command=start_questions, bg="green", fg="white", font=("Arial", 14),
                      width=20)
start_btn.pack(pady=10)

stop_btn = tk.Button(window, text="Stop", command=stop_listening, bg="red", fg="white", font=("Arial", 14),
                     width=20)
stop_btn.pack(pady=5)

evaluate_btn = tk.Button(window, text="Bahola", command=evaluate_answers, bg="blue", fg="white", font=("Arial", 14),
                         width=20)
evaluate_btn.pack(pady=10)

# Savollar va Javoblar
tk.Label(window, text="Savollar:", font=("Arial", 12, "bold"), bg="#f0f0f0").pack(pady=5)
questions_box = tk.Text(window, height=5, width=60, font=("Arial", 11))
questions_box.pack()

tk.Label(window, text="Javoblar:", font=("Arial", 12, "bold"), bg="#f0f0f0").pack(pady=5)
answers_box = tk.Text(window, height=8, width=60, font=("Arial", 11))
answers_box.pack()

window.mainloop()
